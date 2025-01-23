import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from io import BytesIO
import re

FONT_MAP = {
    'Bengali': 'SutonnyMJ',
    'English': 'Times New Roman',
    'Arabic': 'Al Majeed Quranic',
    'Mixed': 'Kalpurus',
    'Nirmala': 'Kalpurus'
}

def fix_bengali_conjuncts(text):
    conjunct_fixes = {
        '\u09CD\u09AF': '\u09CD\u09AF',
        '\u09A8\u09CD\u09A4': '\u09A8\u09CD\u09A4',
        '\u09B7\u09CD\u09A0': '\u09B7\u09CD\u09A0'
    }
    for broken, fixed in conjunct_fixes.items():
        text = text.replace(broken, fixed)
    return text

def remove_watermark(doc):
    for section in doc.sections:
        if hasattr(section._sectPr, 'background'):
            section._sectPr.remove(section._sectPr.background)

def remove_borders(section):
    sectPr = section._sectPr
    pgBorders = sectPr.find(qn('w:pgBorders'))
    if pgBorders is not None:
        sectPr.remove(pgBorders)

def set_single_column(section):
    cols = section._sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'), '1')

def set_run_font(run, font_name, font_size):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.color.rgb = RGBColor(0, 0, 0)

def convert_table_to_image(table):
    table_content = "\n".join(["\t".join([cell.text for cell in row.cells]) for row in table.rows])
    font = ImageFont.load_default()
    width, height = 600, 200
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((10, 10), table_content, fill="black", font=font)
    
    image_stream = BytesIO()
    image.save(image_stream, format="PNG")
    image_stream.seek(0)
    return image_stream

def process_arabic_text(paragraph):
    contains_arabic = any(re.search(r'[\u0600-\u06FF]', run.text) for run in paragraph.runs)
    is_standalone = not any(re.search(r'[A-Za-z\u0980-\u09FF]', run.text) for run in paragraph.runs)
    
    if contains_arabic and is_standalone:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def format_title(paragraph):
    if paragraph.style.name == 'Title':
        for run in paragraph.runs:
            run.font.size = Pt(36 if len(paragraph.text) > 50 else 48)
            run.font.bold = True

def process_headings(doc):
    toc_headings = set()
    for para in doc.paragraphs:
        if para.style.name.startswith('TOC'):
            toc_headings.add(para.text.strip())
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if para.text.strip() in toc_headings:
                para.style = doc.styles['Heading 1']
                for run in para.runs:
                    run.font.size = Pt(20 if len(para.text) > 40 else 22)
                    run.font.bold = True
            else:
                para.style = doc.styles['Normal']

def process_word_file(input_path, output_path):
    try:
        doc = Document(input_path)
        remove_watermark(doc)
        
        # Process sections
        for section in doc.sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.orientation = WD_ORIENTATION.PORTRAIT
            section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
            remove_borders(section)
            set_single_column(section)
            
            # Clear headers and footers
            section.header.is_linked_to_previous = False
            section.footer.is_linked_to_previous = False
            if section.header.paragraphs:
                section.header.paragraphs[0].clear()
            if section.footer.paragraphs:
                section.footer.paragraphs[0].clear()

        # Process headings and TOC
        process_headings(doc)
        
        # Process paragraphs
        for para in doc.paragraphs:
            # Remove unwanted content
            price_pattern = r'http[s]?://|www\.|@|(?<!R)okomari|\$|\d+\s*(?:Tk|BDT|USD|টাকা|/-)|[\u09E6-\u09EF]+\s*(?:টাকা|মাত্র)'
            if re.search(price_pattern, para.text, re.IGNORECASE):
                para.clear()
                continue
            
            # Format paragraph
            format_title(para)
            para.paragraph_format.line_spacing = 1.15
            process_arabic_text(para)
            
            # Process runs
            for run in para.runs:
                run.text = fix_bengali_conjuncts(run.text)
                if re.search(r'[\u0980-\u09FF]', run.text):
                    set_run_font(run, FONT_MAP['Bengali' if not para.style.name == 'Mixed' else 'Nirmala'], 16)
                elif re.search(r'[\u0600-\u06FF]', run.text):
                    set_run_font(run, FONT_MAP['Arabic'], 16)
                else:
                    set_run_font(run, FONT_MAP['English' if not para.style.name == 'Mixed' else 'Nirmala'], 16)

        # Process tables
        for table in doc.tables:
            table_width = sum(cell.width for cell in table.rows[0].cells)
            if table_width > Inches(6):
                image_stream = convert_table_to_image(table)
                para = doc.add_paragraph()
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.add_run().add_picture(image_stream, width=Inches(6))
                table._element.getparent().remove(table._element)

        # Process shapes and images
        for shape in doc.inline_shapes:
            if shape.type == 3:  # Text box
                if shape.text_frame.text:
                    para = doc.add_paragraph(shape.text_frame.text)
                    para.style = doc.styles['Normal']
                shape._element.getparent().remove(shape._element)
            elif shape.type == 2:  # Picture
                shape_element = shape._inline.graphic.graphicData.pic.blipFill.blip
                shape_element.set('embedding', 'rId1')
                shape.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        doc.save(output_path)
        
    except Exception as e:
        print(f"Error processing file {input_path}: {e}")

def main():
    input_dir = './data/Ebook/'
    output_dir = './data/Processed_Ebooks/'
    os.makedirs(output_dir, exist_ok=True)
    
    for file_name in os.listdir(input_dir):
        if file_name.endswith('.docx'):
            input_path = os.path.join(input_dir, file_name)
            output_path = os.path.join(output_dir, file_name)
            process_word_file(input_path, output_path)
            print(f"Processed: {file_name}")

if __name__ == "__main__":
    main()